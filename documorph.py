
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile, os, json, io

# -----------------------------------------------------------------------
#  DocuMorph Engine
# -----------------------------------------------------------------------
class DocuMorphEngine:
    def __init__(self, docx_file: str | io.BytesIO | None = None):
        self.document = Document(docx_file) if docx_file else Document()

    # ---------- global formatting -------------------------------------------------
    def set_font(self, font_name, font_size):
        for para in self.document.paragraphs:
            for run in para.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)

    def set_line_spacing(self, spacing):
        for para in self.document.paragraphs:
            para.paragraph_format.line_spacing = spacing

    def set_alignment(self, alignment):
        align_map = {"Left": WD_ALIGN_PARAGRAPH.LEFT,
                     "Center": WD_ALIGN_PARAGRAPH.CENTER,
                     "Right": WD_ALIGN_PARAGRAPH.RIGHT,
                     "Justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
        for para in self.document.paragraphs:
            para.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

    def set_margins(self, top, bottom, left, right):
        sec = self.document.sections[0]
        sec.top_margin, sec.bottom_margin = Inches(top), Inches(bottom)
        sec.left_margin, sec.right_margin = Inches(left), Inches(right)

    # ---------- logo --------------------------------------------------------------
    def add_logo(self, image_file: io.BytesIO | str, width, height):
        """
        Insert a logo at the very top of the first section header.
        Works with either a Streamlit-uploaded BytesIO or a file path.
        """
        hdr = self.document.sections[0].header
        hdr.is_linked_to_previous = False

        # Ensure we have a paragraph to insert before
        first_para = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        logo_para = first_para.insert_paragraph_before()
        run = logo_para.add_run()

        # rewind BytesIO if needed
        if isinstance(image_file, (io.BytesIO, io.BufferedReader)):
            image_file.seek(0)
        run.add_picture(image_file, width=Inches(width), height=Inches(height))

    # ---------- header / footer ---------------------------------------------------
    def set_header_footer(self, h_text, f_text, size, align):
        align_map = {"Left": WD_ALIGN_PARAGRAPH.LEFT,
                     "Center": WD_ALIGN_PARAGRAPH.CENTER,
                     "Right": WD_ALIGN_PARAGRAPH.RIGHT}
        for sec in self.document.sections:
            hdr, ftr = sec.header, sec.footer
            if not hdr.paragraphs:
                hdr.add_paragraph()
            if not ftr.paragraphs:
                ftr.add_paragraph()

            h_para = hdr.paragraphs[-1]   # after logo, if present
            f_para = ftr.paragraphs[0]

            h_para.text, f_para.text = h_text, f_text
            for p in (h_para, f_para):
                if p.runs:
                    p.runs[0].font.size = Pt(size)
                p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)

    # ---------- content helpers ---------------------------------------------------
    def add_section_title(self, title):        self.document.add_heading(title, level=1)
    def add_bullet_list(self, items):          [self.document.add_paragraph(i, style="List Bullet") for i in items]

    def add_figure(self, image, w, h, caption="", pos="Below"):
        if pos == "Above" and caption:
            self.document.add_paragraph(caption, style="Caption")
        p = self.document.add_paragraph()
        p.add_run().add_picture(image, width=Inches(w), height=Inches(h))
        if pos == "Below" and caption:
            self.document.add_paragraph(caption, style="Caption")

    # ---------- save --------------------------------------------------------------
    def save(self, path): self.document.save(path)

# -----------------------------------------------------------------------
#  Template manager helpers
# -----------------------------------------------------------------------
TEMPLATE_DIR = "templates"
os.makedirs(TEMPLATE_DIR, exist_ok=True)

def list_templates():  return [f[:-5] for f in os.listdir(TEMPLATE_DIR) if f.endswith(".json")]
def load_template(name):  return json.load(open(os.path.join(TEMPLATE_DIR, f"{name}.json"))) if name else {}
def save_template(name, cfg):  json.dump(cfg, open(os.path.join(TEMPLATE_DIR, f"{name}.json"), "w"))
def delete_template(name):  os.remove(os.path.join(TEMPLATE_DIR, f"{name}.json"))

# -----------------------------------------------------------------------
#  Streamlit UI
# -----------------------------------------------------------------------
st.set_page_config("DocuMorph AI", layout="wide")
st.title("📄 DocuMorph AI – Intelligent Formatter")

st.markdown("""
<style>
.stButton>button       {width:100%; padding:0.75em;}
.big-download button   {background:#4E79A7; color:white;}
</style>
""", unsafe_allow_html=True)

# ----- Sidebar : templates -------------------------------------------------------
with st.sidebar:
    st.header("💾 Template Manager")
    sel = st.selectbox("Load template", ["<none>"] + list_templates())
    cfg = load_template(sel if sel != "<none>" else "")

    new_name = st.text_input("Save current as")
    if st.button("💾 Save Template"):
        save_template(new_name, st.session_state)
        st.success(f"Template '{new_name}' saved!")

    if sel != "<none>" and st.button("🗑 Delete Template"):
        delete_template(sel)
        st.warning(f"Template '{sel}' deleted!")
        st.experimental_rerun()

# ----- Tabs ----------------------------------------------------------------------
tabs = st.tabs(["Styling", "Logo & HF", "Figures & Sections", "Export"])

# --- Styling tab -----------------------------------------------------------------
with tabs[0]:
    st.subheader("🎨 Document Styling")
    c1, c2 = st.columns(2)

    with c1:
        font_name   = st.selectbox("Font Style", ["Times New Roman", "Arial", "Calibri", "Georgia"], 0)
        font_size   = st.slider("Font Size", 8, 24, 12)
        line_spacing= st.slider("Line Spacing", 1.0, 2.0, 1.15, 0.05)
    with c2:
        alignment   = st.radio("Alignment", ["Left", "Center", "Right", "Justify"], horizontal=True)
        margins = [st.number_input(lbl, 0.1, 3.0, 1.0, 0.1, key=lbl)
                   for lbl in ("Top Margin","Bottom Margin","Left Margin","Right Margin")]
        st.session_state["margins"] = margins

# --- Logo / Header & Footer tab ---------------------------------------------------
with tabs[1]:
    st.subheader("🖼 Logo & Header/Footer")
    col1, col2 = st.columns(2)
    with col1:
        logo   = st.file_uploader("Upload Logo", type=["png","jpg","jpeg"])
        logo_w = st.slider("Logo Width (in)", 0.5, 4.0, 1.0, 0.1)
        logo_h = st.slider("Logo Height (in)",0.5, 4.0, 1.0, 0.1)
    with col2:
        header_text = st.text_input("Header Text")
        footer_text = st.text_input("Footer Text")
        hf_size     = st.slider("HF Font Size", 8, 20, 10)
        hf_align    = st.selectbox("HF Alignment", ["Left","Center","Right"], 1)

# --- Figures / sections tab -------------------------------------------------------
with tabs[2]:
    st.subheader("📑 Figures & Sections")
    section_title = st.text_input("Section Title")
    bullets_input = st.text_area("Bullet List (one per line)")
    figure   = st.file_uploader("Insert Figure", type=["png","jpg","jpeg"], key="fig")
    fig_w    = st.slider("Figure Width (in)", 1.0, 6.0, 4.0)
    fig_h    = st.slider("Figure Height (in)",1.0, 6.0, 3.0)
    caption      = st.text_input("Caption")
    caption_pos  = st.radio("Caption Position", ["Above","Below"], horizontal=True)

# --- Export tab ------------------------------------------------------------------
with tabs[3]:
    st.subheader("📤 Generate & Download")
    uploaded_file = st.file_uploader("Upload base DOCX", type=["docx"])

    if st.button("📝 Generate & Download"):
        if not uploaded_file:
            st.error("Please upload a DOCX first!")
        else:
            eng = DocuMorphEngine(uploaded_file)
            eng.set_font(font_name, font_size)
            eng.set_line_spacing(line_spacing)
            eng.set_alignment(alignment)
            eng.set_margins(*st.session_state["margins"])

            if logo:  eng.add_logo(logo, logo_w, logo_h)
            eng.set_header_footer(header_text, footer_text, hf_size, hf_align)

            if section_title.strip(): eng.add_section_title(section_title.strip())
            bullets = [b.strip() for b in bullets_input.splitlines() if b.strip()]
            if bullets: eng.add_bullet_list(bullets)
            if figure:  figure.seek(0); eng.add_figure(figure, fig_w, fig_h, caption, caption_pos)

            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            eng.save(tmp.name); tmp.close()
            with open(tmp.name, "rb") as f:
                st.download_button("⬇ Download Document", f, "formatted.docx",
                                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   use_container_width=True, key="download", help="Download the formatted DOCX")
            os.unlink(tmp.name)
